import pandas as pd
import os
from docx import Document
import streamlit as st
from werkzeug.utils import secure_filename

def categorize_students(input_file, outstanding_min, good_min):
    try:
        df = pd.read_excel(input_file, engine='openpyxl')
        
        # Normalize column names
        df.columns = df.columns.str.strip().str.lower()

        # Validate required columns
        required_columns = {'roll number', 'name', 'marks'}
        if not required_columns.issubset(set(df.columns)):
            st.error("Input file must contain 'Roll Number', 'Name', and 'Marks' columns.")
            return None, None

        def categorize(marks):
            if marks >= outstanding_min:
                return 'Advanced Learner'
            elif marks >= good_min:
                return 'Average Learner'
            else:
                return 'Slow Learner'

        df['category'] = df['marks'].apply(categorize)
        output_xlsx = "categorized_students.xlsx"
        df[['roll number', 'name', 'marks', 'category']].to_excel(output_xlsx, index=False)

        # Create Word Document
        output_docx = "categorized_students.docx"
        doc = Document()
        doc.add_heading('Classified Student Roll Numbers', level=1)

        categories = {'Advanced Learner': [], 'Average Learner': [], 'Slow Learner': []}
        for _, row in df.iterrows():
            categories[row['category']].append(str(row['roll number']))

        for category, roll_numbers in categories.items():
            doc.add_heading(category, level=2)
            doc.add_paragraph(", ".join(roll_numbers))

        doc.save(output_docx)
        return output_xlsx, output_docx

    except Exception as e:
        st.error(f"An error occurred: {e}")
        return None, None

# Streamlit UI with Background Image
st.markdown(
    """
    <style>
    body, .stApp {
        background-image: url('https://raw.githubusercontent.com/KiranGaikwad2020/Academic-Automation/Dev/version1.1/mmit-logo.jpg');
        background-size: contain;
        background-position: center;
        background-repeat: no-repeat;
    }
    </style>
    """,
    unsafe_allow_html=True
)

st.title("Student Data Classification Web App by Dr KPG")

uploaded_file = st.file_uploader("Upload Excel File", type=["xlsx"])

outstanding_min = st.number_input("Minimum Marks for Advanced Learner", min_value=0, step=1)
good_min = st.number_input("Minimum Marks for Average Learner", min_value=0, step=1)

if st.button("Categorize Students"):
    if uploaded_file is not None:
        filename = secure_filename(uploaded_file.name)
        with open(filename, "wb") as f:
            f.write(uploaded_file.getbuffer())
        
        xlsx_path, docx_path = categorize_students(filename, outstanding_min, good_min)
        
        if xlsx_path and docx_path:
            st.success("Files generated successfully!")
            st.download_button(label="Download Excel File", data=open(xlsx_path, "rb").read(), file_name=xlsx_path)
            st.download_button(label="Download Word File", data=open(docx_path, "rb").read(), file_name=docx_path)
        
        # Cleanup generated files
        if os.path.exists(filename):
            os.remove(filename)
        if xlsx_path and os.path.exists(xlsx_path):
            os.remove(xlsx_path)
        if docx_path and os.path.exists(docx_path):
            os.remove(docx_path)
    else:
        st.warning("Please upload an Excel file.")
