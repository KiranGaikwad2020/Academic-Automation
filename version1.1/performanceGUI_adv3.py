import tkinter as tk
from tkinter import filedialog, messagebox
from PIL import Image, ImageTk  # Pillow for image processing
import pandas as pd
from docx import Document

# Function to categorize students and generate output files
def categorize_students(input_file, output_xlsx, output_docx, fast_min, average_min):
    try:
        df = pd.read_excel(input_file, engine='openpyxl')

        if 'Roll Number' not in df.columns or 'Name' not in df.columns or 'Marks' not in df.columns:
            messagebox.showerror("Error", "Input file must contain 'Roll Number', 'Name', and 'Marks' columns.")
            return
        
        def categorize(marks):
            if marks >= fast_min:
                return 'Fast Learner'
            elif marks >= average_min:
                return 'Average Learner'
            else:
                return 'Slow Learner'
        
        df['Category'] = df['Marks'].apply(categorize)
        df[['Roll Number', 'Name', 'Marks', 'Category']].to_excel(output_xlsx, index=False)
        
        # Create Word Document
        doc = Document()
        doc.add_heading('Classified Student Roll Numbers', level=1)
        
        categories = {'Fast Learner': [], 'Average Learner': [], 'Slow Learner': []}
        for _, row in df.iterrows():
            categories[row['Category']].append(str(row['Roll Number']))
        
        for category, roll_numbers in categories.items():
            doc.add_heading(category, level=2)
            doc.add_paragraph(", ".join(roll_numbers))
        
        doc.save(output_docx)
        
        messagebox.showinfo("Success", f"Categorized student list saved to:\n{output_xlsx}\nand\n{output_docx}")
    except Exception as e:
        messagebox.showerror("Error", f"An error occurred: {e}")

# Function to open file dialog
def select_input_file():
    file_path = filedialog.askopenfilename(filetypes=[("Excel Files", "*.xlsx;*.xls"), ("All Files", "*.*")])
    entry_input_file.delete(0, tk.END)
    entry_input_file.insert(0, file_path)

# Function to start processing
def start_categorization():
    input_file = entry_input_file.get()
    output_xlsx = "categorized_students.xlsx"
    output_docx = "categorized_students.docx"
    try:
        fast_min = int(entry_fast.get())
        average_min = int(entry_average.get())

        if not input_file:
            messagebox.showwarning("Warning", "Please select an input file.")
            return
        
        categorize_students(input_file, output_xlsx, output_docx, fast_min, average_min)
    
    except ValueError:
        messagebox.showerror("Error", "Please enter valid numerical values for marks.")

# Create GUI window
root = tk.Tk()
root.title("Student Data Classification")
root.geometry("500x400")
root.resizable(False, False)

# Load background image
bg_image = Image.open("mmit-logo.jpg").resize((500, 400), Image.LANCZOS)
bg_photo = ImageTk.PhotoImage(bg_image)
canvas = tk.Canvas(root, width=500, height=400)
canvas.pack(fill="both", expand=True)
canvas.create_image(0, 0, image=bg_photo, anchor="nw")

# Labels and Entry Widgets
canvas.create_text(250, 30, text="Student Data Classification", font=("Arial", 14, "bold"), fill="black")

tk.Label(root, text="Select Input Excel File:", font=("Arial", 11, "bold"), bg="white").place(x=40, y=60)
entry_input_file = tk.Entry(root, width=40, relief="solid")
entry_input_file.place(x=40, y=85)
tk.Button(root, text="Browse", command=select_input_file, bg="#008000", fg="#ffffff").place(x=400, y=80)

tk.Label(root, text="Minimum Marks for Fast Learner:", font=("Arial", 11, "bold"), bg="white").place(x=40, y=120)
entry_fast = tk.Entry(root, width=10, relief="solid")
entry_fast.place(x=320, y=120)

tk.Label(root, text="Minimum Marks for Average Learner:", font=("Arial", 11, "bold"), bg="white").place(x=40, y=160)
entry_average = tk.Entry(root, width=10, relief="solid")
entry_average.place(x=320, y=160)

# Process Button
tk.Button(root, text="Categorize Students", command=start_categorization, bg="#008000", fg="#ffffff", font=("Arial", 12, "bold")).place(x=170, y=220)

# Run the GUI
root.mainloop()

